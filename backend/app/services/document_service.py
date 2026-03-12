"""
Document Processing Service
Handles PDF, Word, Excel, and Images (with OCR)
Chunks documents for RAG system
"""

import os
import uuid
import hashlib
import asyncio
import fitz  # PyMuPDF for PDF
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from pdf2image import convert_from_path  # TAMBAHKAN INI
from docx import Document as DocxDocument

# Document loaders
import fitz  # PyMuPDF for PDF
from docx import Document as DocxDocument
import openpyxl
from PIL import Image
import pytesseract

# Text processing
from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.core.database import add_documents, delete_documents, get_collection
from app.services.llm_service import get_ollama_service


class DocumentProcessor:
    """Process various document types for RAG"""
    
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.processed_dir = Path(settings.PROCESSED_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        
        # Text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
    
    def _generate_doc_id(self, content: str, filename: str) -> str:
        """Generate unique document ID"""
        hash_input = f"{filename}:{content[:100]}:{datetime.now().isoformat()}"
        return hashlib.md5(hash_input.encode()).hexdigest()[:16]

    async def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file (dengan OCR fallback untuk dokumen hasil scan)"""
        text = ""
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                page_text = page.get_text()

                # Jika ada teks digital biasa
                if page_text and page_text.strip():
                    text += f"\n--- Halaman {page_num + 1} ---\n{page_text}"
                else:
                    # OCR Fallback: Jika halaman berupa gambar/hasil scan
                    try:
                        # Convert halaman PDF ke gambar
                        images = convert_from_path(file_path, first_page=page_num+1, last_page=page_num+1)
                        for img in images:
                            # Gunakan Tesseract untuk membaca gambar (bahasa Indo + Inggris)
                            ocr_text = pytesseract.image_to_string(img, lang='ind+eng')
                            if ocr_text.strip():
                                text += f"\n--- Halaman {page_num + 1} (OCR) ---\n{ocr_text}"
                    except Exception as ocr_err:
                        print(f"Error OCR Halaman {page_num + 1}: {ocr_err}")
            doc.close()
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        return text.strip()
    
    async def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        text = ""
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract from tables too
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        return text.strip()
    
    async def extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        text = ""
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"\n--- Sheet: {sheet_name} ---\n"
                
                for row in sheet.iter_rows():
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value))
                    if row_values:
                        text += " | ".join(row_values) + "\n"
            wb.close()
        except Exception as e:
            print(f"Error extracting Excel: {e}")
        return text.strip()
    
    async def extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR (Tesseract)"""
        text = ""
        try:
            image = Image.open(file_path)
            # OCR with Indonesian + English language
            text = pytesseract.image_to_string(
                image, 
                lang='ind+eng',
                config='--psm 6'  # Assume uniform block of text
            )
        except Exception as e:
            print(f"Error OCR: {e}")
        return text.strip()
    
    async def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading text file: {e}")
            return ""
    
    async def process_document(
        self, 
        file_path: str, 
        filename: str,
        metadata: Dict[str, Any] = None
    ) -> Tuple[bool, str, int]:
        """
        Process a document and add to vector database
        Returns: (success, message, chunk_count)
        """
        ext = Path(filename).suffix.lower()
        
        # Extract text based on file type
        if ext == ".pdf":
            text = await self.extract_text_from_pdf(file_path)
        elif ext in [".doc", ".docx"]:
            text = await self.extract_text_from_docx(file_path)
        elif ext in [".xls", ".xlsx"]:
            text = await self.extract_text_from_excel(file_path)
        elif ext in [".jpg", ".jpeg", ".png", ".gif", ".webp"]:
            text = await self.extract_text_from_image(file_path)
        elif ext in [".txt", ".md"]:
            text = await self.extract_text_from_txt(file_path)
        else:
            return False, f"Format file tidak didukung: {ext}", 0
        
        if not text or len(text.strip()) < 10:
            return False, "Tidak dapat mengekstrak teks dari dokumen", 0
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        if not chunks:
            return False, "Dokumen tidak menghasilkan chunk", 0
        
        # Generate embeddings
        ollama = get_ollama_service()
        embeddings = await ollama.get_embeddings_batch(chunks)
        
        if not embeddings or len(embeddings) != len(chunks):
            return False, "Gagal membuat embeddings", 0
        
        # Prepare metadata
        base_metadata = {
            "filename": filename,
            "file_type": ext,
            "upload_date": datetime.now().isoformat(),
            "chunk_count": len(chunks),
            **(metadata or {})
        }
        
        # Prepare documents for database
        doc_ids = []
        doc_metadatas = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = f"{self._generate_doc_id(chunk, filename)}_{i}"
            doc_ids.append(chunk_id)
            doc_metadatas.append({
                **base_metadata,
                "chunk_index": i,
                "chunk_id": chunk_id
            })
        
        # Add to vector database
        success = await add_documents(
            documents=chunks,
            embeddings=embeddings,
            metadatas=doc_metadatas,
            ids=doc_ids
        )
        
        if success:
            return True, f"Berhasil memproses {len(chunks)} chunk", len(chunks)
        else:
            return False, "Gagal menyimpan ke database", 0
    
    async def search_similar(
        self, 
        query: str, 
        top_k: int = None,
        filter_metadata: Dict = None
    ) -> List[Dict[str, Any]]:
        """Search for similar documents"""
        from app.core.database import query_documents
        
        ollama = get_ollama_service()
        query_embedding = await ollama.get_embedding(query)
        
        if not query_embedding:
            return []
        
        results = await query_documents(
            query_embedding=query_embedding,
            n_results=top_k or settings.TOP_K_RESULTS,
            where=filter_metadata
        )
        
        # Format results
        similar_docs = []
        documents = results.get("documents", [[]])[0]
        metadatas = results.get("metadatas", [[]])[0]
        distances = results.get("distances", [[]])[0]
        
        for doc, meta, dist in zip(documents, metadatas, distances):
            # Convert distance to similarity score (cosine distance to similarity)
            similarity = 1 - dist
            
            if similarity >= settings.SIMILARITY_THRESHOLD:
                similar_docs.append({
                    "content": doc,
                    "metadata": meta,
                    "similarity": round(similarity, 4)
                })
        
        return similar_docs
    
    async def delete_document_by_filename(self, filename: str) -> Tuple[bool, str]:
        """Delete all chunks of a document by filename"""
        try:
            collection = get_collection()
            
            # Get all chunks with this filename
            results = collection.get(
                where={"filename": filename},
                include=["metadatas"]
            )
            
            if not results["ids"]:
                return False, "Dokumen tidak ditemukan"
            
            # Delete all chunks
            await delete_documents(results["ids"])
            
            return True, f"Berhasil menghapus {len(results['ids'])} chunk"
        except Exception as e:
            return False, f"Error: {str(e)}"
    
    async def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about stored documents"""
        try:
            collection = get_collection()
            total_count = collection.count()
            
            # Get all documents to analyze
            all_docs = collection.get(include=["metadatas"])
            
            # Count unique files
            unique_files = set()
            file_types = {}
            
            for meta in all_docs.get("metadatas", []):
                if meta:
                    filename = meta.get("filename", "unknown")
                    file_type = meta.get("file_type", "unknown")
                    unique_files.add(filename)
                    file_types[file_type] = file_types.get(file_type, 0) + 1
            
            return {
                "total_chunks": total_count,
                "unique_documents": len(unique_files),
                "file_types": file_types,
                "documents": list(unique_files)
            }
        except Exception as e:
            return {"error": str(e)}


# Global instance
document_processor = DocumentProcessor()


def get_document_processor() -> DocumentProcessor:
    """Get document processor instance"""
    return document_processor
